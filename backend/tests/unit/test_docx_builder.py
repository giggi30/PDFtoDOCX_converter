import io
from collections.abc import Iterable

from docx import Document
from docx.oxml.ns import qn

from app.conversion.docx_builder import document_model_to_docx
from app.conversion.models import (
    BoundingBox,
    Decoration,
    DocumentModel,
    ImageElement,
    PageModel,
    Region,
    TextBlock,
    TextStyle,
)

_ONE_PIXEL_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z0iUAAAAASUVORK5CYII="
)


def _style(font: str = "Helvetica", *, bold: bool = False) -> TextStyle:
    return TextStyle(
        font_family=font,
        font_size=11,
        color="#123456",
        bold=bold,
    )


def _block(
    block_id: str,
    text: str,
    x0: float,
    y0: float,
    x1: float,
    *,
    block_type: str = "paragraph",
    style: TextStyle | None = None,
) -> TextBlock:
    return TextBlock.model_validate(
        {
            "id": block_id,
            "bbox": {"x0": x0, "y0": y0, "x1": x1, "y1": y0 + 14},
            "text": text,
            "style": style or _style(),
            "block_type": block_type,
        }
    )


def _document_text(document: object) -> str:
    values: Iterable[object] = document.element.body.iter(qn("w:t"))
    return "\n".join(str(value.text) for value in values if value.text)


def test_builds_editable_single_column_docx_with_styles_and_image() -> None:
    region = Region(
        type="main",
        bbox=BoundingBox(x0=48, y0=40, x1=564, y1=740),
        blocks=[
            _block(
                "heading",
                "Anonymous Candidate",
                48,
                40,
                300,
                block_type="heading",
                style=_style(bold=True),
            ),
            _block("item", "• Editable content", 48, 100, 300, block_type="list_item"),
        ],
    )
    model = DocumentModel(
        source_type="native",
        pages=[
            PageModel(
                page_number=1,
                width_pt=612,
                height_pt=792,
                regions=[region],
                images=[
                    ImageElement(
                        id="portrait",
                        bbox=BoundingBox(x0=420, y0=40, x1=492, y1=112),
                        width_px=1,
                        height_px=1,
                        role="photo",
                        mime_type="image/png",
                        content_base64=_ONE_PIXEL_PNG,
                    )
                ],
            )
        ],
    )

    result = document_model_to_docx(model)
    document = Document(io.BytesIO(result.content))

    assert "Anonymous Candidate" in _document_text(document)
    assert "Editable content" in _document_text(document)
    assert "• Editable content" not in _document_text(document)
    assert document.sections[0].page_width.pt == 612
    assert document.sections[0].page_height.pt == 792
    assert document.paragraphs[0].runs[0].font.name == "Arial"
    assert str(document.paragraphs[0].runs[0].font.color.rgb) == "123456"
    assert len(document.inline_shapes) == 1
    assert not result.warnings


def test_builds_sidebar_as_borderless_shaded_table() -> None:
    sidebar = Region(
        type="sidebar",
        bbox=BoundingBox(x0=24, y0=40, x1=180, y1=740),
        blocks=[_block("contact", "CONTACT", 24, 40, 160, block_type="heading")],
    )
    main = Region(
        type="main",
        bbox=BoundingBox(x0=214, y0=40, x1=564, y1=740),
        blocks=[_block("experience", "EXPERIENCE", 214, 40, 500, block_type="heading")],
    )
    model = DocumentModel(
        source_type="native",
        pages=[
            PageModel(
                page_number=1,
                width_pt=612,
                height_pt=792,
                regions=[sidebar, main],
                decorations=[
                    Decoration(
                        id="sidebar-background",
                        bbox=BoundingBox(x0=0, y0=0, x1=180, y1=792),
                        kind="rectangle",
                        fill_color="#DDEEFF",
                    )
                ],
            )
        ],
    )

    result = document_model_to_docx(model)
    document = Document(io.BytesIO(result.content))
    table = document.tables[0]

    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "CONTACT"
    assert table.cell(0, 1).text == "EXPERIENCE"
    shading = table.cell(0, 0)._tc.tcPr.find(qn("w:shd"))
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert shading is not None
    assert shading.get(qn("w:fill")) == "DDEEFF"
    assert borders is not None
    assert borders.find(qn("w:top")).get(qn("w:val")) == "nil"
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    assert table_width is not None
    assert table_width.get(qn("w:type")) == "dxa"
    assert not result.warnings


def test_reports_substitutions_and_elements_that_cannot_be_reproduced() -> None:
    region = Region(
        type="main",
        bbox=BoundingBox(x0=48, y0=40, x1=564, y1=740),
        blocks=[_block("body", "Text", 48, 40, 300, style=_style("UnknownFont"))],
    )
    model = DocumentModel(
        source_type="native",
        pages=[
            PageModel(
                page_number=1,
                width_pt=612,
                height_pt=792,
                regions=[region],
                images=[
                    ImageElement(
                        id="missing-image",
                        bbox=BoundingBox(x0=48, y0=80, x1=100, y1=132),
                    )
                ],
                decorations=[
                    Decoration(
                        id="line",
                        bbox=BoundingBox(x0=48, y0=150, x1=300, y1=150),
                        kind="line",
                    )
                ],
            )
        ],
    )

    result = document_model_to_docx(model)

    assert any("UnknownFont" in warning for warning in result.warnings)
    assert any("missing-image" in warning for warning in result.warnings)
    assert any("'line'" in warning for warning in result.warnings)


def test_creates_one_word_section_for_each_source_page() -> None:
    pages = []
    for page_number, width in ((1, 612), (2, 595)):
        region = Region(
            type="main",
            bbox=BoundingBox(x0=48, y0=40, x1=width - 48, y1=740),
            blocks=[_block(f"page-{page_number}", f"Page {page_number}", 48, 40, 200)],
        )
        pages.append(
            PageModel(
                page_number=page_number,
                width_pt=width,
                height_pt=792,
                regions=[region],
            )
        )

    result = document_model_to_docx(DocumentModel(source_type="native", pages=pages))
    document = Document(io.BytesIO(result.content))

    assert len(document.sections) == 2
    assert [section.page_width.pt for section in document.sections] == [612, 595]
    assert "Page 1" in _document_text(document)
    assert "Page 2" in _document_text(document)


def test_hero_title_keeps_source_font_and_drops_shadow_duplicates() -> None:
    shadow = TextBlock.model_validate(
        {
            "id": "shadow-title",
            "bbox": {"x0": 72, "y0": 42, "x1": 330, "y1": 70},
            "text": "LUIGI RAGNI",
            "style": {"font_family": "EBGaramond-Regular", "font_size": 22, "color": "#0B1F33"},
            "block_type": "heading",
        }
    )
    title = TextBlock.model_validate(
        {
            "id": "title",
            "bbox": {"x0": 72, "y0": 40, "x1": 330, "y1": 68},
            "text": "LUIGI RAGNI",
            "style": {"font_family": "EBGaramond-Regular", "font_size": 22, "color": "#FFFFFF"},
            "block_type": "heading",
        }
    )
    hero = Region(
        type="hero",
        bbox=BoundingBox(x0=0, y0=0, x1=612, y1=150),
        blocks=[shadow, title],
    )
    model = DocumentModel(
        source_type="native",
        pages=[
            PageModel(
                page_number=1,
                width_pt=612,
                height_pt=792,
                regions=[hero],
                decorations=[
                    Decoration(
                        id="hero-background",
                        bbox=hero.bbox,
                        kind="rectangle",
                        fill_color="#082F64",
                    )
                ],
            )
        ],
    )

    result = document_model_to_docx(model)
    document = Document(io.BytesIO(result.content))

    assert document.element.body.xpath("count(.//w:t[normalize-space(.)='LUIGI RAGNI'])") == 1
    title_paragraph = next(
        paragraph for paragraph in document.tables[0].cell(0, 0).paragraphs if paragraph.text.strip()
    )
    assert title_paragraph.runs[0].font.name == "EB Garamond"
    assert not result.warnings


def test_builds_zoned_canva_layout_and_distinguishes_high_fidelity_mode() -> None:
    hero = Region(
        type="hero",
        bbox=BoundingBox(x0=0, y0=0, x1=612, y1=150),
        blocks=[
            _block(
                "title",
                "Anonymous Candidate",
                72,
                40,
                300,
                block_type="heading",
                style=TextStyle(
                    font_family="EBGaramond-Regular",
                    font_size=22,
                    color="#FFFFFF",
                ),
            )
        ],
    )
    main = Region(
        type="main",
        bbox=BoundingBox(x0=48, y0=180, x1=564, y1=700),
        blocks=[
            _block(
                "profile",
                "PROFILE",
                48,
                180,
                150,
                block_type="heading",
                style=_style("OpenSans-Bold", bold=True),
            ),
            _block("summary", "Editable summary", 48, 215, 400),
            _block(
                "skill",
                "• Source spacing",
                48,
                250,
                260,
                block_type="list_item",
            ),
        ],
    )
    card = Region(
        type="card",
        bbox=BoundingBox(x0=310, y0=650, x1=600, y1=780),
        blocks=[_block("contact", "candidate@example.test", 340, 680, 560)],
    )
    model = DocumentModel(
        source_type="native",
        pages=[
            PageModel(
                page_number=1,
                width_pt=612,
                height_pt=792,
                regions=[hero, main, card],
                images=[
                    ImageElement(
                        id="portrait",
                        bbox=BoundingBox(x0=450, y0=0, x1=600, y1=150),
                        width_px=1,
                        height_px=1,
                        role="photo",
                        mime_type="image/png",
                        content_base64=_ONE_PIXEL_PNG,
                    )
                ],
                decorations=[
                    Decoration(
                        id="hero-background",
                        bbox=hero.bbox,
                        kind="rectangle",
                        fill_color="#082F64",
                    ),
                    Decoration(
                        id="card-background",
                        bbox=card.bbox,
                        kind="rectangle",
                        fill_color="#DEECFF",
                    ),
                    Decoration(
                        id="profile-accent",
                        bbox=BoundingBox(x0=48, y0=196, x1=130, y1=196),
                        kind="line",
                        stroke_color="#082F64",
                    ),
                ],
            )
        ],
    )

    result = document_model_to_docx(model)
    doc = Document(io.BytesIO(result.content))

    assert len(doc.tables) == 2
    # The hero background is an absolute page anchor instead of cell shading,
    # so it cannot spill below the source header when rendered by LibreOffice.
    assert doc.tables[0].cell(0, 0)._tc.tcPr.find(qn("w:shd")) is None
    doc_width = doc.tables[0]._tbl.tblPr.find(qn("w:tblW"))
    assert doc_width.get(qn("w:w")) == str(612 * 20)
    list_paragraph = next(
        paragraph for paragraph in doc.paragraphs if paragraph.style.name == "List Bullet"
    )
    assert len(doc.element.xpath(".//wp:anchor")) >= 3
    contextual_spacing = list_paragraph._p.pPr.find(qn("w:contextualSpacing"))
    assert contextual_spacing is not None
    assert contextual_spacing.get(qn("w:val")) == "0"
    assert not result.warnings
