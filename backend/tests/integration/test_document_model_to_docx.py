import io
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.conversion.docx_builder import document_model_to_docx
from app.conversion.pipeline import pdf_to_document_model
from tests.fixtures.pdf_factory import CV_FIXTURES, CvFixture

_PROJECT_ROOT = Path(__file__).resolve().parents[3]


@pytest.mark.parametrize("fixture", CV_FIXTURES, ids=lambda fixture: fixture.name)
def test_ten_cv_models_produce_editable_docx_with_matching_text(fixture: CvFixture) -> None:
    model = pdf_to_document_model(fixture.pdf)

    result = document_model_to_docx(model)
    document = Document(io.BytesIO(result.content))
    extracted_text = "\n".join(
        node.text for node in document.element.body.iter(qn("w:t")) if node.text
    )
    header_text = "\n".join(
        paragraph.text for section in document.sections for paragraph in section.header.paragraphs
    )
    extracted_text = f"{header_text}\n{extracted_text}"

    assert all(expected in extracted_text for expected in fixture.expected_text)
    if fixture.layout in {"sidebar", "columns"}:
        assert len(document.tables) == 1
        assert len(document.tables[0].columns) == 2


def test_canva_fixture_builds_full_bleed_two_column_docx() -> None:
    source = _PROJECT_ROOT / "test-documents/fixtures/Curriculum Vitae.pdf"
    model = pdf_to_document_model(source.read_bytes())

    result = document_model_to_docx(model)
    document = Document(io.BytesIO(result.content))

    assert not result.warnings
    assert len(document.tables) == 2
    assert len(document.tables[1].columns) == 2
    assert document.sections[0].left_margin.pt == 0
    assert document.sections[0].right_margin.pt == 0
    assert len(document.element.xpath(".//wp:anchor")) >= 3
    right_border = document.tables[1].cell(0, 0)._tc.tcPr.find(
        qn("w:tcBorders")
    ).find(qn("w:right"))
    assert right_border is not None
    assert right_border.get(qn("w:val")) == "single"
    extracted_text = "\n".join(
        node.text for node in document.element.body.iter(qn("w:t")) if node.text
    )
    assert "Carlo Giddi" in extracted_text
    assert "Personal Trainer" in extracted_text
