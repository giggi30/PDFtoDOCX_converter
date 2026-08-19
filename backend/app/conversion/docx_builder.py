import base64
import binascii
import io
import re
from copy import deepcopy
from dataclasses import dataclass
from typing import cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Emu, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image, UnidentifiedImageError

from app.conversion.models import (
    BoundingBox,
    Decoration,
    DocumentModel,
    ImageElement,
    PageModel,
    Region,
    TextBlock,
    TextStyle,
)

_LIST_MARKER = re.compile(r"^\s*(?:[•‣▪]|[-–—])\s+")
_FONT_MAPPINGS = {
    "arial": "Arial",
    "courier": "Courier New",
    "couriernew": "Courier New",
    "clearsans": "Arial",
    "ebgaramond": "EB Garamond",
    "ebgaramondregular": "EB Garamond",
    "garamond": "EB Garamond",
    "eb": "EB Garamond",
    "helvetica": "Arial",
    "leaguespartan": "Arial",
    "opensans": "Open Sans",
    "times": "Times New Roman",
    "timesnewroman": "Times New Roman",
}

_TITLE_STYLE = "CV Title"
_HEADING_STYLE = "CV Heading 1"


@dataclass(frozen=True)
class DocxBuildResult:
    content: bytes
    warnings: tuple[str, ...]


def _normalised_font_key(value: str) -> str:
    return re.sub(r"[^a-z]", "", value.lower()).replace("bold", "").replace("italic", "")


def _mapped_font(style: TextStyle, warnings: list[str]) -> str:
    if style.font_family is None:
        return "Arial"
    key = _normalised_font_key(style.font_family)
    for source, target in _FONT_MAPPINGS.items():
        if source in key:
            return target
    warnings.append(f"Font '{style.font_family}' is unsupported; Arial was substituted.")
    return "Arial"


def _set_run_font(run: Run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:cs"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)


def _apply_text_style(paragraph: Paragraph, block: TextBlock, warnings: list[str]) -> None:
    text = _LIST_MARKER.sub("", block.text) if block.block_type == "list_item" else block.text
    lines = text.split("\n")
    run = paragraph.add_run(lines[0] if lines else "")
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)
    font_name = _mapped_font(block.style, warnings)
    _set_run_font(run, font_name)
    run.font.size = Pt(max(6.0, min(block.style.font_size, 48.0)))
    run.font.bold = block.style.bold or block.block_type == "heading"
    run.font.italic = block.style.italic
    run.font.underline = block.style.underline
    run.font.color.rgb = RGBColor.from_string(block.style.color.removeprefix("#"))


def _configure_paragraph(paragraph: Paragraph, block: TextBlock) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4 if block.block_type == "heading" else 2)
    source_lines = max(1, len(block.text.splitlines()))
    source_line_height = block.bbox.height / source_lines
    paragraph.paragraph_format.line_spacing = max(
        1.0,
        min(1.3, source_line_height / block.style.font_size),
    )
    paragraph.paragraph_format.keep_with_next = block.block_type == "heading"
    if block.block_type == "heading":
        paragraph.style = _TITLE_STYLE if block.style.font_size >= 16 else _HEADING_STYLE
    elif block.block_type == "list_item":
        paragraph.style = "List Bullet"
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.first_line_indent = Pt(-12)
        # The built-in Word style enables contextual spacing, which makes
        # LibreOffice discard source-derived gaps between consecutive bullets.
        # An explicit paragraph override keeps each item's original vertical
        # rhythm when the DOCX is rendered outside Microsoft Word.
        contextual_spacing = OxmlElement("w:contextualSpacing")
        contextual_spacing.set(qn("w:val"), "0")
        paragraph._p.get_or_add_pPr().append(contextual_spacing)


def _new_paragraph(container: DocumentObject | _Cell, first: bool) -> Paragraph:
    if isinstance(container, _Cell) and first:
        paragraph = container.paragraphs[0]
        paragraph.clear()
        return cast(Paragraph, paragraph)
    return container.add_paragraph()


def _add_image(
    container: DocumentObject | _Cell,
    image: ImageElement,
    max_width_pt: float,
    first: bool,
    warnings: list[str],
    max_height_pt: float | None = None,
    space_before_pt: float = 0,
) -> None:
    if not image.content_base64:
        warnings.append(
            f"Image '{image.id}' could not be embedded because its content is unavailable."
        )
        return
    try:
        content = base64.b64decode(image.content_base64, validate=True)
        paragraph = _new_paragraph(container, first)
        paragraph.paragraph_format.space_before = Pt(space_before_pt)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        width_pt = max(1.0, min(image.bbox.width, max_width_pt))
        height_pt = max(1.0, image.bbox.height * width_pt / max(image.bbox.width, 1.0))
        if max_height_pt is not None and height_pt > max_height_pt:
            width_pt *= max_height_pt / height_pt
            height_pt = max_height_pt
        run.add_picture(io.BytesIO(content), width=Pt(width_pt), height=Pt(height_pt))
    except (binascii.Error, ValueError, TypeError, UnrecognizedImageError):
        warnings.append(f"Image '{image.id}' contains invalid data and was skipped.")


def _inline_to_page_anchor(
    inline: object,
    x_pt: float,
    y_pt: float,
    *,
    behind_document: bool = False,
) -> None:
    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "0" if behind_document else "251658240",
        "behindDoc": "1" if behind_document else "0",
        "locked": "0",
        "layoutInCell": "0",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)

    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "page")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = str(round(Emu(Pt(x_pt))))
    horizontal.append(horizontal_offset)
    anchor.append(horizontal)

    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "page")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = str(round(Emu(Pt(y_pt))))
    vertical.append(vertical_offset)
    anchor.append(vertical)

    for tag in ("wp:extent", "wp:effectExtent"):
        child = inline.find(qn(tag))  # type: ignore[attr-defined]
        if child is not None:
            anchor.append(deepcopy(child))
    anchor.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        child = inline.find(qn(tag))  # type: ignore[attr-defined]
        if child is not None:
            anchor.append(deepcopy(child))
    inline.getparent().replace(inline, anchor)  # type: ignore[attr-defined]


def _add_floating_image(
    container: DocumentObject | _Cell | Paragraph,
    page: PageModel,
    image: ImageElement,
    first: bool,
    warnings: list[str],
    *,
    behind_document: bool = False,
) -> None:
    if not image.content_base64:
        warnings.append(
            f"Image '{image.id}' could not be embedded because its content is unavailable."
        )
        return
    try:
        content = base64.b64decode(image.content_base64, validate=True)
        visible_x0 = max(0.0, image.bbox.x0)
        visible_y0 = max(0.0, image.bbox.y0)
        visible_x1 = min(page.width_pt, image.bbox.x1)
        visible_y1 = min(page.height_pt, image.bbox.y1)
        if visible_x1 <= visible_x0 or visible_y1 <= visible_y0:
            return
        needs_crop = (
            visible_x0 != image.bbox.x0
            or visible_y0 != image.bbox.y0
            or visible_x1 != image.bbox.x1
            or visible_y1 != image.bbox.y1
        )
        if needs_crop:
            source: Image.Image = Image.open(io.BytesIO(content))
            source.load()
            left = round((visible_x0 - image.bbox.x0) / image.bbox.width * source.width)
            top = round((visible_y0 - image.bbox.y0) / image.bbox.height * source.height)
            right = round((visible_x1 - image.bbox.x0) / image.bbox.width * source.width)
            bottom = round((visible_y1 - image.bbox.y0) / image.bbox.height * source.height)
            source = source.crop((left, top, right, bottom))
            stream = io.BytesIO()
            source.save(stream, format="PNG")
            content = stream.getvalue()
        if isinstance(container, Paragraph):
            paragraph = container
        else:
            paragraph = _new_paragraph(container, first)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(1)
        run = paragraph.add_run()
        run.font.size = Pt(1)
        shape = run.add_picture(
            io.BytesIO(content),
            width=Pt(max(1.0, visible_x1 - visible_x0)),
            height=Pt(max(1.0, visible_y1 - visible_y0)),
        )
        _inline_to_page_anchor(
            shape._inline,
            visible_x0,
            visible_y0,
            behind_document=behind_document,
        )
    except (
        binascii.Error,
        UnidentifiedImageError,
        ValueError,
        TypeError,
        UnrecognizedImageError,
    ):
        warnings.append(f"Image '{image.id}' contains invalid data and was skipped.")


def _add_solid_band(
    container: DocumentObject | _Cell | Paragraph,
    page: PageModel,
    bbox: BoundingBox,
    color: str,
    first: bool,
) -> None:
    visible_x0 = max(0.0, bbox.x0)
    visible_y0 = max(0.0, bbox.y0)
    visible_x1 = min(page.width_pt, bbox.x1)
    visible_y1 = min(page.height_pt, bbox.y1)
    if visible_x1 <= visible_x0 or visible_y1 <= visible_y0:
        return
    red, green, blue = (
        int(color[index : index + 2], 16) for index in (1, 3, 5)
    )
    source = Image.new("RGB", (1, 1), (red, green, blue))
    stream = io.BytesIO()
    source.save(stream, format="PNG")
    if isinstance(container, Paragraph):
        paragraph = container
    else:
        paragraph = _new_paragraph(container, first)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
    run = paragraph.add_run()
    run.font.size = Pt(1)
    shape = run.add_picture(
        io.BytesIO(stream.getvalue()),
        width=Pt(visible_x1 - visible_x0),
        height=Pt(visible_y1 - visible_y0),
    )
    _inline_to_page_anchor(
        shape._inline,
        visible_x0,
        visible_y0,
        behind_document=True,
    )


def _add_region_content(
    container: DocumentObject | _Cell,
    region: Region,
    images: list[ImageElement],
    max_width_pt: float,
    warnings: list[str],
    *,
    preserve_positions: bool = False,
    decorations: list[Decoration] | None = None,
    left_indent_pt: float = 0,
    right_indent_pt: float = 0,
    gap_scale: float = 0.98,
    max_gap_pt: float = 18.0,
    preserve_horizontal_positions: bool = False,
) -> None:
    # Separate icons that should be rendered inline with specific text blocks
    icon_map: dict[str, ImageElement] = {}
    standalone_images: list[ImageElement] = []
    for image in images:
        if image.role == "icon" or (image.bbox.width <= 36 and image.bbox.height <= 36):
            paired_block = next(
                (
                    b
                    for b in region.blocks
                    if abs((b.bbox.y0 + b.bbox.y1) / 2 - (image.bbox.y0 + image.bbox.y1) / 2) <= 14
                    and image.bbox.x0 - 8 <= b.bbox.x0
                ),
                None,
            )
            if paired_block is not None:
                icon_map[paired_block.id] = image
                continue
        standalone_images.append(image)

    # Deduplicate blocks to handle Canva's shadow text layers
    region_blocks = _deduplicate_shadow_blocks(region.blocks)
    _ = decorations
    
    elements: list[tuple[float, float, TextBlock | ImageElement]] = [
        (block.bbox.y0, block.bbox.x0, block) for block in region_blocks
    ]
    elements.extend(
        (image.bbox.y0, image.bbox.x0, image) 
        for image in standalone_images 
    )
    first = True
    previous_bottom: float | None = region.bbox.y0 if preserve_positions else None
    # Use the provided gap scale to better match PDF visual spacing in Word
    active_gap_scale = gap_scale
    
    for _, _, element in sorted(elements, key=lambda item: (item[0], item[1])):
        space_before = 0.0
        if preserve_positions and previous_bottom is not None:
            gap = max(0.0, element.bbox.y0 - previous_bottom)
            # Word adds a small internal padding to paragraphs; subtracting 1pt 
            # helps prevent cumulative downward drift and overlapping.
            space_before = min(max_gap_pt, max(0.0, gap * active_gap_scale - 1.0))

        if isinstance(element, ImageElement):
            _add_image(
                container, 
                element, 
                max_width_pt, 
                first, 
                warnings, 
                space_before_pt=space_before
            )
        else:
            paragraph = _new_paragraph(container, first)
            _configure_paragraph(paragraph, element)
            paragraph.paragraph_format.space_before = Pt(space_before)
            
            if left_indent_pt > 0:
                curr = paragraph.paragraph_format.left_indent
                base_indent = curr.pt if curr else 0
                paragraph.paragraph_format.left_indent = Pt(base_indent + left_indent_pt)
            if right_indent_pt > 0:
                curr = paragraph.paragraph_format.right_indent
                base_indent = curr.pt if curr else 0
                paragraph.paragraph_format.right_indent = Pt(base_indent + right_indent_pt)
            
            if preserve_horizontal_positions:
                rel_x = max(0.0, element.bbox.x0 - region.bbox.x0)
                curr = paragraph.paragraph_format.left_indent
                paragraph.paragraph_format.left_indent = Pt((curr.pt if curr else 0) + rel_x)
            
            paired_icon = icon_map.get(element.id)

            if paired_icon is not None and paired_icon.content_base64:
                try:
                    icon_content = base64.b64decode(paired_icon.content_base64)
                    run_icon = paragraph.add_run()
                    max_icon_h = element.style.font_size * 1.35
                    icon_size = max(10.0, min(paired_icon.bbox.height, max_icon_h))
                    scale = icon_size / max(paired_icon.bbox.height, 1.0)
                    icon_w = paired_icon.bbox.width * scale
                    run_icon.add_picture(
                        io.BytesIO(icon_content),
                        width=Pt(icon_w),
                        height=Pt(icon_size),
                    )
                    space_run = paragraph.add_run("  ")
                    _set_run_font(space_run, _mapped_font(element.style, warnings))
                except Exception:
                    pass
            _apply_text_style(paragraph, element, warnings)
        
        previous_bottom = element.bbox.y1
        first = False


def _set_cell_shading(cell: _Cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color.removeprefix("#"))


def _set_cell_margins(
    cell: _Cell,
    value: int = 120,
    *,
    top: int | None = None,
    start: int | None = None,
    bottom: int | None = None,
    end: int | None = None,
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    values = {
        "top": top,
        "start": start,
        "left": start,
        "bottom": bottom,
        "end": end,
        "right": end,
    }
    for edge in ("top", "start", "left", "bottom", "end", "right"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value if values[edge] is None else values[edge]))
        node.set(qn("w:type"), "dxa")


def _configure_layout_table(
    table: Table,
    widths_pt: list[float],
    minimum_height_pt: float,
    *,
    indent_pt: float = 0,
    exact_height: bool = False,
) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(round(sum(widths_pt) * 20)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(round(indent_pt * 20)))
    table_indent.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_pt, strict=True):
        grid_col.w = Pt(width)
    for cell, width in zip(table.rows[0].cells, widths_pt, strict=True):
        cell.width = Pt(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(cell)
    table.rows[0].height = Pt(minimum_height_pt)
    table.rows[0].height_rule = (
        WD_ROW_HEIGHT_RULE.EXACTLY if exact_height else WD_ROW_HEIGHT_RULE.AT_LEAST
    )


def _is_shadow_duplicate(block: TextBlock, other: TextBlock) -> bool:
    text1 = block.text.strip().lower()
    text2 = other.text.strip().lower()
    if not text1 or not text2:
        return False

    # Relax text match: if one is a subset of the other, or if they overlap spatially
    # and have similar font sizes, they are likely duplicates/shadows.
    text_match = text1 == text2 or text1 in text2 or text2 in text1
    
    # Calculate intersection of bounding boxes
    x_overlap = max(0, min(block.bbox.x1, other.bbox.x1) - max(block.bbox.x0, other.bbox.x0))
    y_overlap = max(0, min(block.bbox.y1, other.bbox.y1) - max(block.bbox.y0, other.bbox.y0))
    area_overlap = x_overlap * y_overlap
    
    significant_overlap = area_overlap > 0 and (
        area_overlap / (block.bbox.width * block.bbox.height) > 0.5 or
        area_overlap / (other.bbox.width * other.bbox.height) > 0.5
    )

    if not (text_match or significant_overlap):
        return False

    # Check for proximity even if no direct area overlap
    horizontal_match = (
        abs(block.bbox.x0 - other.bbox.x0) < 20
        or abs(block.bbox.x1 - other.bbox.x1) < 20
    )
    vertical_match = abs(block.bbox.y0 - other.bbox.y0) < 12

    if not (horizontal_match and vertical_match):
        return False

    if abs(block.style.font_size - other.style.font_size) > 3.0:
        return False

    return block.bbox.height > 0 and other.bbox.height > 0


def _deduplicate_shadow_blocks(blocks: list[TextBlock]) -> list[TextBlock]:
    kept: list[TextBlock] = []
    for block in sorted(blocks, key=lambda item: (item.bbox.y0, item.bbox.x0)):
        if any(_is_shadow_duplicate(block, candidate) for candidate in kept):
            continue
        kept.append(block)
    return kept


def _configure_document_styles(document: DocumentObject) -> None:
    # Set default font for Normal style
    normal_font = document.styles["Normal"].font
    normal_font.name = "Arial"
    
    for name, size, outline_level in (
        (_TITLE_STYLE, 22, 0),
        (_HEADING_STYLE, 14, 1),
    ):
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.size = Pt(size)
        style.font.bold = True
        # Try to use EB Garamond as default for headings if possible
        style.font.name = "EB Garamond"
        properties = style.element.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(outline_level))
        properties.append(outline)


def _images_for_region(page: PageModel, region: Region) -> list[ImageElement]:
    result: list[ImageElement] = []
    for image in page.images:
        if image.role == "background":
            continue
        centre_x = (image.bbox.x0 + image.bbox.x1) / 2
        centre_y = (image.bbox.y0 + image.bbox.y1) / 2
        inside_x = region.bbox.x0 - 20 <= centre_x <= region.bbox.x1 + 20
        inside_y = region.bbox.y0 - 20 <= centre_y <= region.bbox.y1 + 20
        if inside_x and (inside_y or region.type not in {"hero", "card"}):
            result.append(image)
    return result


def _region_fill(page: PageModel, region: Region) -> str | None:
    preferred: str | None = None
    for decoration in page.decorations:
        if decoration.kind != "rectangle" or decoration.fill_color is None:
            continue
        horizontal_overlap = max(
            0.0,
            min(decoration.bbox.x1, region.bbox.x1) - max(decoration.bbox.x0, region.bbox.x0),
        )
        vertical_overlap = max(
            0.0,
            min(decoration.bbox.y1, region.bbox.y1) - max(decoration.bbox.y0, region.bbox.y0),
        )
        if (
            horizontal_overlap >= region.bbox.width * 0.75
            and vertical_overlap >= region.bbox.height * 0.75
        ):
            if decoration.fill_color not in {"#000000", "#FFFFFF"}:
                return decoration.fill_color
            preferred = preferred or decoration.fill_color
    return preferred


def _sidebar_fill(page: PageModel, region: Region) -> str | None:
    return _region_fill(page, region)


def _configure_section(section: Section, page: PageModel) -> None:
    section.page_width = Pt(page.width_pt)
    section.page_height = Pt(page.height_pt)
    has_hero = any(region.type == "hero" for region in page.regions)
    has_zones = any(region.type in {"hero", "card"} for region in page.regions)
    section.top_margin = Pt(0 if has_hero else 12)
    section.bottom_margin = Pt(0 if has_zones else 12)
    section.left_margin = Pt(0 if has_zones else 36)
    section.right_margin = Pt(0 if has_zones else 36)
    section.header_distance = Pt(0 if has_hero else 18)
    section.footer_distance = Pt(18)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def _add_header_or_footer(
    section: Section, _page: PageModel, region: Region, warnings: list[str]
) -> None:
    target = section.header if region.type == "header" else section.footer
    target.is_linked_to_previous = False
    paragraph = target.paragraphs[0]
    paragraph.clear()
    if region.type == "footer":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, block in enumerate(region.blocks):
        current = paragraph if index == 0 else target.add_paragraph()
        _configure_paragraph(current, block)
        _apply_text_style(current, block, warnings)


def _add_spacer(document: DocumentObject, height_pt: float) -> None:
    if height_pt <= 0:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(height_pt)
    paragraph.paragraph_format.line_spacing = Pt(1)
    run = paragraph.add_run("")
    run.font.size = Pt(1)


def _add_hero(
    document: DocumentObject,
    page: PageModel,
    hero: Region,
    warnings: list[str],
) -> tuple[set[str], float]:
    hero_blocks = _deduplicate_shadow_blocks(hero.blocks)
    images = _images_for_region(page, hero)
    band_top = max(0.0, hero.bbox.y0)
    band_bottom = min(page.height_pt, hero.bbox.y1)
    hero_height = max(1.0, band_bottom - band_top)
    table = document.add_table(rows=1, cols=1)
    _configure_layout_table(
        table,
        [page.width_pt],
        hero_height,
        indent_pt=0,
        exact_height=True,
    )
    fill = _region_fill(page, hero) or "#082F64"
    cell = table.cell(0, 0)
    _set_cell_margins(cell, 0)

    used: set[str] = set()
    floating_p = _new_paragraph(cell, True)
    floating_p.paragraph_format.space_before = Pt(0)
    floating_p.paragraph_format.space_after = Pt(0)
    floating_p.paragraph_format.line_spacing = Pt(1)
    # Keep the hero fill as a page-anchored rectangle only.  Cell shading is
    # flow-driven in LibreOffice: when the name paragraph changes the row's
    # effective height, the shading can extend below the source header on one
    # side.  The anchored band remains clipped to the original bounding box.
    _add_solid_band(floating_p, page, hero.bbox, fill, first=False)
    for background in (image for image in page.images if image.role == "background"):
        _add_floating_image(
            floating_p,
            page,
            background,
            first=False,
            warnings=warnings,
            behind_document=True,
        )
    footer = _footer_band(page)
    if footer is not None:
        _add_solid_band(
            floating_p,
            page,
            footer.bbox,
            footer.fill_color or "#45545F",
            first=False,
        )
    for image in images:
        _add_floating_image(floating_p, page, image, first=False, warnings=warnings)
        used.add(image.id)

    band_blocks = [
        block
        for block in hero_blocks
        if (block.bbox.y0 + block.bbox.y1) / 2 <= band_bottom + 2
    ]
    outside_blocks = [block for block in hero_blocks if block not in band_blocks]
    previous_bottom = band_top
    for block in sorted(band_blocks, key=lambda item: (item.bbox.y0, item.bbox.x0)):
        paragraph = _new_paragraph(cell, False)
        _configure_paragraph(paragraph, block)
        paragraph.paragraph_format.left_indent = Pt(max(0.0, block.bbox.x0))
        target_right = max(block.bbox.x1 + 8, block.bbox.x0 + block.bbox.width * 1.25)
        desired_right = min(page.width_pt, target_right)
        paragraph.paragraph_format.right_indent = Pt(
            max(0.0, page.width_pt - desired_right)
        )
        gap = max(0.0, block.bbox.y0 - previous_bottom)
        max_allowed_before = max(0.0, hero_height - block.style.font_size - 24)
        if previous_bottom == band_top:
            space_before = min(max_allowed_before, max(0.0, gap - 55.0))
        else:
            space_before = gap * 0.85
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(max(10.0, block.style.font_size * 0.9))
        _apply_text_style(paragraph, block, warnings)
        previous_bottom = block.bbox.y1

    flow_bottom = band_bottom
    for block in sorted(outside_blocks, key=lambda item: (item.bbox.y0, item.bbox.x0)):
        paragraph = document.add_paragraph()
        _configure_paragraph(paragraph, block)
        paragraph.paragraph_format.left_indent = Pt(max(0.0, block.bbox.x0))
        target_right = max(block.bbox.x1 + 8, block.bbox.x0 + block.bbox.width * 1.25)
        desired_right = min(page.width_pt, target_right)
        paragraph.paragraph_format.right_indent = Pt(
            max(0.0, page.width_pt - desired_right)
        )
        paragraph.paragraph_format.space_before = Pt(
            max(0.0, block.bbox.y0 - flow_bottom) * 0.8
        )
        paragraph.paragraph_format.space_after = Pt(0)
        _apply_text_style(paragraph, block, warnings)
        flow_bottom = block.bbox.y1
    return used, flow_bottom


def _add_card_zone(
    document: DocumentObject,
    page: PageModel,
    main: Region,
    card: Region,
    warnings: list[str],
) -> None:
    before = [block for block in main.blocks if block.bbox.y0 < card.bbox.y0 - 2]
    alongside = [block for block in main.blocks if block not in before]
    if before:
        before_region = main.model_copy(update={"blocks": before})
        _add_region_content(
            document,
            before_region,
            [],
            page.width_pt - 72,
            warnings,
            preserve_positions=True,
            decorations=page.decorations,
            left_indent_pt=36,
            right_indent_pt=36,
            gap_scale=0.7,
            max_gap_pt=30,
        )
    last_bottom = max((block.bbox.y1 for block in before), default=card.bbox.y0)
    _add_spacer(document, min(32.0, max(0.0, card.bbox.y0 - last_bottom) * 0.45))

    table_width = page.width_pt
    table_start = 0.0
    left_width = min(max(card.bbox.x0 - table_start, table_width * 0.35), table_width * 0.68)
    widths = [left_width, table_width - left_width]
    card_fill = _region_fill(page, card) or "#DEECFF"
    _add_solid_band(document, page, card.bbox, card_fill, first=False)
    table = document.add_table(rows=1, cols=2)
    # The visual rectangle is an absolute background anchor. Keeping the flow
    # table at its natural height prevents a one-row table from being moved
    # wholesale to the next page when the card overlaps the main column.
    _configure_layout_table(
        table,
        widths,
        1.0,
        indent_pt=0,
        exact_height=False,
    )

    left_margin = max(2.0, 39.0 - table_start)
    _set_cell_margins(
        table.cell(0, 0),
        0,
        top=120,
        start=0,
        bottom=60,
        end=100,
    )
    side_region = main.model_copy(update={"blocks": alongside})
    _add_region_content(
        table.cell(0, 0),
        side_region,
        [],
        max(1.0, widths[0] - left_margin - 5),
        warnings,
        preserve_positions=True,
        decorations=page.decorations,
        left_indent_pt=left_margin,
    )

    card_text_x = min((block.bbox.x0 for block in card.blocks), default=card.bbox.x0 + 28)
    card_icon_x = min(
        (image.bbox.x0 for image in _images_for_region(page, card) if image.role == "icon"),
        default=card_text_x,
    )
    # Start at the icon's source position, not at the text's.  This keeps the
    # text width available after the icon and preserves the final contact row.
    card_margin = max(2.0, min(card_text_x, card_icon_x) - card.bbox.x0)
    _set_cell_margins(
        table.cell(0, 1),
        0,
        top=440,
        start=0,
        bottom=80,
        end=80,
    )
    _add_region_content(
        table.cell(0, 1),
        card,
        _images_for_region(page, card),
        max(1.0, widths[1] - card_margin - 4),
        warnings,
        preserve_positions=True,
        left_indent_pt=card_margin,
    )


def _set_cell_border(cell: _Cell, edge: str, color: str, size: int = 8) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color.removeprefix("#"))


def _vertical_divider(page: PageModel) -> Decoration | None:
    return next(
        (
            decoration
            for decoration in page.decorations
            if decoration.kind in {"line", "rectangle"}
            and decoration.bbox.width <= 2
            and decoration.bbox.height >= page.height_pt * 0.3
            and page.width_pt * 0.2 <= decoration.bbox.x0 <= page.width_pt * 0.8
        ),
        None,
    )


def _footer_band(page: PageModel) -> Decoration | None:
    return next(
        (
            decoration
            for decoration in page.decorations
            if decoration.kind == "rectangle"
            and decoration.fill_color not in {None, "#000000", "#FFFFFF"}
            and decoration.bbox.width >= page.width_pt * 0.75
            and decoration.bbox.y0 >= page.height_pt * 0.8
        ),
        None,
    )


def _add_zoned_columns(
    document: DocumentObject,
    page: PageModel,
    columns: list[Region],
    warnings: list[str],
    used_images: set[str] | None = None,
) -> set[str]:
    ordered = sorted(columns[:2], key=lambda r: r.bbox.x0)
    divider = _vertical_divider(page)
    if divider is not None:
        boundary = (divider.bbox.x0 + divider.bbox.x1) / 2
    else:
        boundary = (ordered[0].bbox.x1 + ordered[1].bbox.x0) / 2
    boundary = min(max(boundary, page.width_pt * 0.25), page.width_pt * 0.75)
    widths = [boundary, page.width_pt - boundary]

    table = document.add_table(rows=1, cols=2)
    _configure_layout_table(
        table,
        widths,
        1.0,
        exact_height=False,
    )
    if divider is not None:
        _set_cell_border(
            table.cell(0, 0),
            "right",
            divider.stroke_color or divider.fill_color or "#45545F",
            size=max(4, round((divider.stroke_width or 0.7) * 8)),
        )

    cell_starts = [0.0, boundary]
    cell_ends = [boundary, page.width_pt]
    columns_used_images: set[str] = set()
    for cell, region, width, cell_start, cell_end in zip(
        table.rows[0].cells,
        ordered,
        widths,
        cell_starts,
        cell_ends,
        strict=True,
    ):
        _set_cell_margins(cell, 0)
        left_indent = max(0.0, region.bbox.x0 - cell_start)
        right_indent = max(0.0, cell_end - region.bbox.x1)
        region_images = [
            img for img in _images_for_region(page, region)
            if used_images is None or img.id not in used_images
        ]
        columns_used_images.update(img.id for img in region_images)
        _add_region_content(
            cell,
            region,
            region_images,
            max(1.0, width - left_indent - right_indent),
            warnings,
            preserve_positions=True,
            decorations=page.decorations,
            left_indent_pt=left_indent,
            right_indent_pt=right_indent,
            gap_scale=0.85,
            max_gap_pt=48,
            preserve_horizontal_positions=True,
        )
    return columns_used_images


def _add_zoned_page(
    document: DocumentObject,
    page: PageModel,
    warnings: list[str],
) -> set[str]:
    hero = next((region for region in page.regions if region.type == "hero"), None)
    main = next((region for region in page.regions if region.type == "main"), None)
    card = next((region for region in page.regions if region.type == "card"), None)
    columns = [region for region in page.regions if region.type == "column"]
    used_images: set[str] = set()
    flow_bottom = 0.0
    if hero is not None:
        hero_images, flow_bottom = _add_hero(document, page, hero, warnings)
        used_images.update(hero_images)

    body_regions = list(columns) if columns else ([main] if main is not None else [])
    if card is not None:
        body_regions.append(card)
    if hero is not None and body_regions:
        body_top = min(region.bbox.y0 for region in body_regions)
        has_content_below_band = any(
            (block.bbox.y0 + block.bbox.y1) / 2 > hero.bbox.y1 + 2
            for block in hero.blocks
        )
        gap_scale = 0.9 if has_content_below_band else 0.5
        _add_spacer(document, max(0.0, body_top - flow_bottom) * gap_scale)

    if len(columns) >= 2:
        col_images = _add_zoned_columns(document, page, columns, warnings, used_images)
        used_images.update(col_images)
    elif main is not None and card is not None:
        _add_card_zone(document, page, main, card, warnings)
    elif main is not None:
        images = [image for image in page.images if image.id not in used_images]
        _add_region_content(
            document,
            main,
            images,
            page.width_pt - 72,
            warnings,
            preserve_positions=True,
            decorations=page.decorations,
            left_indent_pt=36,
            right_indent_pt=36,
            gap_scale=0.7,
            max_gap_pt=30,
        )
        used_images.update(image.id for image in images)
    return used_images


def _record_unreproduced_decorations(
    page: PageModel, warnings: list[str],
) -> None:
    lines = [
        decoration
        for decoration in page.decorations
        if decoration.kind == "line"
        and max(decoration.bbox.width, decoration.bbox.height) > 160
    ]
    if any(region.type == "column" for region in page.regions):
        lines = [line for line in lines if line.bbox.height < page.height_pt * 0.3]
    if lines:
        warnings.append(
            f"Pagina {page.page_number}: {len(lines)} linea/e decorativa/e non sono state "
            f"riprodotte esattamente (ad esempio '{lines[0].id}')."
        )


def _add_page(
    document: DocumentObject,
    section: Section,
    page: PageModel,
    warnings: list[str],
) -> None:
    _configure_section(section, page)
    body_regions = [region for region in page.regions if region.type not in {"header", "footer"}]
    for region in page.regions:
        if region.type in {"header", "footer"}:
            _add_header_or_footer(section, page, region, warnings)

    if any(region.type in {"hero", "card"} for region in body_regions):
        used_images = _add_zoned_page(document, page, warnings)
        for image in page.images:
            if image.id not in used_images and not any(
                image in _images_for_region(page, region)
                for region in body_regions
                if region.type in {"hero", "card"}
            ):
                _add_image(document, image, page.width_pt - 72, False, warnings)
    elif len(body_regions) == 1 and body_regions[0].type == "main":
        region = body_regions[0]
        _add_region_content(
            document,
            region,
            page.images,
            page.width_pt - 72,
            warnings,
            preserve_positions=True,
            decorations=page.decorations,
        )
    elif len(body_regions) >= 2:
        ordered = sorted(body_regions[:2], key=lambda item: item.bbox.x0)
        available_width = page.width_pt - 72
        source_widths = [max(region.bbox.width, 1.0) for region in ordered]
        total_source_width = sum(source_widths)
        widths = [available_width * width / total_source_width for width in source_widths]
        table = document.add_table(rows=1, cols=2)
        _configure_layout_table(table, widths, max(1.0, page.height_pt - 84))
        for cell, region, width in zip(table.rows[0].cells, ordered, widths, strict=True):
            _add_region_content(
                cell,
                region,
                _images_for_region(page, region),
                width - 12,
                warnings,
            )
            if region.type == "sidebar":
                fill = _sidebar_fill(page, region)
                if fill:
                    _set_cell_shading(cell, fill)
    elif page.images:
        for image in page.images:
            _add_image(document, image, page.width_pt - 72, False, warnings)

    _record_unreproduced_decorations(page, warnings)


def document_model_to_docx(
    model: DocumentModel,
) -> DocxBuildResult:
    if not isinstance(model, DocumentModel):
        raise TypeError("model must be a DocumentModel")
    document = Document()
    _configure_document_styles(document)
    warnings: list[str] = list(model.warnings)
    if not model.pages:
        document.add_paragraph("")
        warnings.append("The document contains no pages.")
    for index, page in enumerate(model.pages):
        section = document.sections[0] if index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        _add_page(document, section, page, warnings)
    stream = io.BytesIO()
    document.save(stream)
    return DocxBuildResult(content=stream.getvalue(), warnings=tuple(dict.fromkeys(warnings)))
